# Text chunking, extraction
#This module handles all document processing - extracting text from PDFs, DOCX, and TXT files, cleaning the text, 
# and splitting it into manageable chunks for the RAG system.
import re 
from typing import List, Dict, Any
import PyPDF2
import docx
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)  # Replace multiple spaces/newlines with a single space
    text = re.sub(r'[^\w\s.,!?;:()\-"\'\n]', '', text)  # Remove special characters but keep basic punctuation
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Remove multiple newlines

    return text.strip()

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            sentence_ends= []
            for punct in ['.', '!', '?', ';']:
                pos = text.rfind(punct, start, end)
                if pos > start:
                    sentence_ends.append(pos)
            if sentence_ends:
                end = max(sentence_ends) + 1
            else:
                space_pos = text.rfind(' ', start, end)
                if space_pos > start:
                    end = space_pos
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
        
    return chunks

def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n"
                    text += page_text + "\n"
            except Exception as e:
                logger.error(f"Error extracting text from page {page_num + 1}: {e}")
                continue
        if not text.strip():
            raise ValueError("No text can be extracted from PDF.")
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise Exception("Failed to extract text from PDF.")
    
def extract_text_from_docx(file_content: bytes) -> str:
    try :
        doc = docx.Document(BytesIO(file_content))
        text_parts = []
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        if not text_parts:
            raise ValueError("No text can be extracted from DOCX.")
        text = '\n'.join(text_parts)
        return clean_text(text)
    except Exception as e:      
        logger.error(f"Error extracting text from DOCX: {e}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from plain text file"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                return clean_text(text)
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        raise Exception(f"Failed to extract text from TXT file: {str(e)}")

def extract_text_from_txt(file_content: bytes, filename: str) -> dict[str , Any]:
    
    extention = filename.split('.')[-1].lower() if '.' in filename else ''
    try:
        if extention == 'txt':
            text = extract_text_from_txt(file_content)
            doc_type = 'text file'
        elif extention == 'pdf':
            text = extract_text_from_pdf(file_content)
            doc_type = 'pdf file'   
        elif extention == 'docx':
            text = extract_text_from_docx(file_content)
            doc_type = 'docx file'
        else:
            raise ValueError(f"Unsupported file type: {extention}")

        if len(text.strip()) < 50:
            raise ValueError("too short or empty document.")
        
        word_count = len(text.split())
        char_count = len(text)
        
        return{
            'text': text,
            'doc_type': doc_type,
            'word_count': word_count,
            'char_count': char_count,
            'estimated reading time': max(1, word_count // 200)
        }
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        raise Exception(f"Failed to extract text from {filename}: {str(e)}")